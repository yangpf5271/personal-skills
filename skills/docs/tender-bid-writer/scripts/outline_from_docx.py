#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
outline_from_docx.py — 大纲 .docx → outline.json 的确定性提取与转写校验

为什么存在：大纲若靠人工/模型转写成 outline.json，可能被"顺手优化"——
典型事故：转写时按评分表重排章节（把评分因素当章节名）、把子级拉平。
这违反"用户大纲结构（文本/层级/顺序/父子）不可改"的永久规则，且下游
plan.py / generate.py 都以 outline.json 为唯一事实源，入口错则全错。

本脚本提供两道确定性防线：
  1. 提取：直接从源 docx 按标题样式/大纲级别/编号文本解析标题树，绕过手工转写。
  2. 校验：把既有 outline.json 与源 docx 的标题序列逐项比对（层级+文本），
     不一致打印差异并返回码 1，阻止带错进入规划阶段。

用法：
  python outline_from_docx.py 大纲.docx                      # 只打印识别出的标题树
  python outline_from_docx.py 大纲.docx -o outline.json      # 提取（若目标已存在则保留其 meta）
  python outline_from_docx.py 大纲.docx --check outline.json # 校验转写是否与源一致
依赖：pip install python-docx --break-system-packages
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖：pip install python-docx --break-system-packages")
    sys.exit(2)

CN_NUM = r'[一二三四五六七八九十百零〇\d]+'

# 文本编号 → 层级（按优先级匹配；match 后编号前缀会被剥掉）
TEXT_PATTERNS = [
    (re.compile(rf'^第{CN_NUM}[章篇部]\s*[、.．:：]?\s*'), lambda m: 1),
    (re.compile(r'^(\d+(?:[\.．]\d+)+)[\.．]?\s*[、]?\s*'),
     lambda m: len(re.split(r'[\.．]', m.group(1)))),          # 1.1 / 1.1.1 …
    (re.compile(r'^(\d+)\s*[、．.]\s+'), lambda m: 1),           # "1. " / "1、"
    (re.compile(rf'^[一二三四五六七八九十]+\s*、\s*'), lambda m: 1),  # 一、
    (re.compile(rf'^（{CN_NUM}）\s*'), lambda m: 2),             # （一）
    (re.compile(rf'^\({CN_NUM}\)\s*'), lambda m: 2),
]

STYLE_RE = re.compile(r'(?:heading|标题)\s*(\d+)', re.I)
MAX_HEADING_LEN = 60   # 编号文本判定标题时的长度上限，防止把长备注误判为标题


def style_level(p):
    """从段落样式名取标题层级（Heading 1 / 标题 1）。"""
    try:
        m = STYLE_RE.search(p.style.name or "")
        return int(m.group(1)) if m else None
    except Exception:
        return None


def outline_level(p):
    """从段落 XML 的 w:outlineLvl 取层级（0 基 → 1 基）。"""
    try:
        pPr = p._p.pPr
        if pPr is not None:
            el = pPr.find(qn('w:outlineLvl'))
            if el is not None:
                return int(el.get(qn('w:val'))) + 1
    except Exception:
        pass
    return None


def detect(p, styles_only):
    """判定段落是否为标题。返回 (level, 去编号后的标题文本)；非标题返回 (None, 原文本)。
    styles_only=True 时只认标题样式/outlineLvl——文档用了真标题样式时，
    带编号的正文段落（如"（一）xxx"）绝不能靠文本形态猜成标题。"""
    text = (p.text or "").strip()
    if not text:
        return None, ""
    lvl = style_level(p) or outline_level(p)
    if lvl:
        # 样式标题也可能带手打编号，剥掉
        for pat, _fn in TEXT_PATTERNS:
            m = pat.match(text)
            if m:
                text = text[m.end():].strip()
                break
        return (min(lvl, 9), text) if lvl and text else (None, "")
    if styles_only:
        return None, text
    # 全文无样式标题 → 靠文本编号判定（限制长度防误判）
    if len(text) <= MAX_HEADING_LEN:
        for pat, fn in TEXT_PATTERNS:
            m = pat.match(text)
            if m:
                title = text[m.end():].strip()
                if title:
                    return min(fn(m), 9), title
    return None, text


def extract(docx_path):
    """从 docx 提取 [(level, title), ...] 与备注映射 {标题序号: note}。
    文档含 ≥3 个样式标题（Heading/标题 或 outlineLvl）时进入"仅样式"模式，
    文本编号回退只用于完全没有样式标题的纯文本大纲。"""
    doc = Document(docx_path)
    styled = sum(1 for p in doc.paragraphs
                 if (p.text or "").strip() and (style_level(p) or outline_level(p)))
    styles_only = styled >= 3
    items, notes = [], {}
    for p in doc.paragraphs:
        lvl, text = detect(p, styles_only)
        if lvl:
            items.append((lvl, text))
        elif text and items:
            i = len(items) - 1
            notes[i] = (notes.get(i, "") + text).strip()
    mode = f"标题样式（{styled} 个样式标题）" if styles_only else "文本编号回退（未检测到样式标题）"
    print(f"检测模式：{mode}")
    return items, notes


PUNCT_ONLY = re.compile(r'^[…．.、，,；;：:\s—\-–]+$')
ANNOT_SUFFIX = re.compile(r'^(.+?)\s*[（(][^）)]{1,4}[）)]\s*$')


def clean(items, notes, keep_preamble=False, strip_annot=False):
    """清洗提取结果：
    1. 丢弃首个一级标题之前的所有"标题"（大纲 docx 开头常贴排版/装订要求，会被误判）。
    2. 丢弃纯标点/省略号占位标题（如 "……"）。
    3. 对形如 "项目管理（肖）" 的尾部短括注给出警告（疑似分工注记，需人工确认去留）。
    返回 (items, notes, warnings)。"""
    warns = []
    first_h1 = next((i for i, (l, _t) in enumerate(items) if l == 1), 0)
    if first_h1 > 0 and not keep_preamble:
        for _l, t in items[:first_h1]:
            warns.append(f"已丢弃首个一级标题前的疑似非大纲行：{t[:40]}")
        notes = {k - first_h1: v for k, v in notes.items() if k >= first_h1}
        items = items[first_h1:]
    out_items, out_notes, j = [], {}, 0
    for i, (lvl, title) in enumerate(items):
        if PUNCT_ONLY.match(title):
            warns.append(f"已丢弃占位标题：「{title}」（h{lvl}）")
            continue
        m = ANNOT_SUFFIX.match(title)
        if m:
            if strip_annot:
                warns.append(f"已剥离标题尾部括注：「{title}」→「{m.group(1)}」")
                title = m.group(1)
            else:
                warns.append(f"标题「{title}」疑带分工/批注括注，加 --strip-annot 可自动剥离")
        if i in notes:
            out_notes[j] = notes[i]
        out_items.append((lvl, title))
        j += 1
    return out_items, out_notes, warns


def build_tree(items, notes):
    """[(level,title)] → 递归 outline 树。层级跳跃时挂到最近的浅层父节点下。"""
    root = {"children": []}
    stack = [(0, root)]   # (level, node)
    for i, (lvl, title) in enumerate(items):
        while stack and stack[-1][0] >= lvl:
            stack.pop()
        parent = stack[-1][1] if stack else root
        node = {"title": title}
        if i in notes:
            node["note"] = notes[i]
        parent.setdefault("children", []).append(node)
        stack.append((lvl, node))
    return root["children"]


def flatten_json(outline, depth=1):
    out = []
    for nd in outline:
        out.append((depth, (nd.get("title") or "").strip()))
        if nd.get("children"):
            out.extend(flatten_json(nd["children"], depth + 1))
    return out


def norm(t):
    """比较用归一化：去所有空白，全角括号转半角。"""
    t = "".join((t or "").split())
    return t.replace("（", "(").replace("）", ")")


def print_tree(items, notes=None):
    for i, (lvl, title) in enumerate(items):
        note = f"   [note: {notes[i][:30]}…]" if notes and i in notes else ""
        print(f"{'  ' * (lvl - 1)}h{lvl} {title}{note}")


def do_check(src_items, json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    got = flatten_json(data.get("outline", []))
    exp = [(l, t) for (l, t) in src_items]
    bad = []
    def eq(e, g):
        if e[0] != g[0]:
            return False
        if norm(e[1]) == norm(g[1]):
            return True
        # 仅差尾部短括注（分工注记）视为一致：源「项目管理（肖）」= outline「项目管理」
        m = ANNOT_SUFFIX.match(e[1])
        return bool(m and norm(m.group(1)) == norm(g[1]))

    for i in range(max(len(exp), len(got))):
        e = exp[i] if i < len(exp) else None
        g = got[i] if i < len(got) else None
        if e and g and eq(e, g):
            continue
        bad.append((i, e, g))
    if not bad:
        print(f"校验通过：{len(exp)} 个标题与源大纲逐级一致。")
        return 0
    print(f"校验失败：{len(bad)} 处不一致（源 {len(exp)} 项 / outline.json {len(got)} 项）：")
    for i, e, g in bad[:20]:
        es = f"h{e[0]} {e[1]}" if e else "（源已结束）"
        gs = f"h{g[0]} {g[1]}" if g else "（outline.json 已结束）"
        print(f"  #{i+1:>3} 源大纲: {es}\n       outline: {gs}")
    if len(bad) > 20:
        print(f"  …… 另有 {len(bad) - 20} 处")
    print("\n结构以源大纲为准。禁止按评分表重排章节/拉平子级；请修正 outline.json 后重跑校验。")
    return 1


def main():
    ap = argparse.ArgumentParser(description="大纲 docx → outline.json 提取与转写校验")
    ap.add_argument("docx", help="源大纲 .docx")
    ap.add_argument("-o", "--out", help="输出 outline.json（已存在则保留其 meta）")
    ap.add_argument("--check", metavar="OUTLINE_JSON", help="校验既有 outline.json 与源大纲是否一致")
    ap.add_argument("--keep-preamble", action="store_true",
                    help="保留首个一级标题之前识别出的标题（默认丢弃并警告）")
    ap.add_argument("--strip-annot", action="store_true",
                    help="自动剥离标题尾部的短括注（如 分工注记（某某））")
    args = ap.parse_args()

    items, notes = extract(args.docx)
    if not items:
        print("未在文档中识别出任何标题：请确认大纲使用了标题样式或 1 / 1.1 / 一、 等编号。")
        return 2
    items, notes, warns = clean(items, notes, keep_preamble=args.keep_preamble,
                                strip_annot=args.strip_annot)
    for w in warns:
        print(f"[警告] {w}")
    if warns:
        print()
    if not items:
        print("清洗后无标题剩余，请检查源文档。")
        return 2

    if args.check:
        return do_check(items, args.check)

    print(f"识别出 {len(items)} 个标题：\n")
    print_tree(items, notes)

    if args.out:
        meta = {}
        if os.path.exists(args.out):
            try:
                with open(args.out, "r", encoding="utf-8") as f:
                    meta = json.load(f).get("meta", {}) or {}
            except Exception:
                pass
        data = {"meta": meta, "outline": build_tree(items, notes)}
        with open(args.out, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"\n已写入 {args.out}（meta 请补充/核对）。请肉眼比对上方标题树与源文档。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
