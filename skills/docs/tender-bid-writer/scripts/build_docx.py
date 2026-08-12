#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — 投标文件排版引擎 (Tender bid document formatting engine)

输入: 一个结构化 JSON 文件 (content model)，描述封面 + 正文块。
输出: 一个完全按中国政府/招标投标排版规范格式化的 .docx。

引擎负责"确定性"的排版工作，不负责写内容。它自动完成：
  - A4 纸 / 页边距 (上2.5cm 其余2cm)
  - 五级标题样式 (三号宋体加粗居中，段前段后1行，一级标题另起一页)
  - 标题自动编号 1 / 1.1 / 1.1.1 / 1.1.1.1 / 1.1.1.1.1
  - 正文 (宋体/Times New Roman 小四 1.5倍行距 首行缩进2字符 两端对齐)
  - 三级列表序号 （1） / 1） / ①  (小四宋体不加粗)
  - 全框线表格，表题在上(五号宋体)，表头加粗，内容居中，分章编号 表X-Y
  - 图片居中无缩进，图题在下(五号宋体加粗)，分章编号 图X-Y
  - 封面、目录(自动域)、正文 三段式分节
  - 页码: 正文从1开始阿拉伯数字，页脚居中，Times New Roman 五号；封面目录不编号
  - 自动目录域 (TOC \\o "1-3")，打开时提示更新

用法:
  python build_docx.py content.json output.docx
  (content.json 的字段说明见 references/content_schema.md)
"""

import sys
import json
import os
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# ============================ 排版配置 ============================
# 所有排版参数集中到 format_config.json（见 scripts/format_config.json）。
# 优先级：--format-config 指定 > content.json 同目录 > 本脚本同目录 > 下面的内置默认。
# JSON 缺失或字段不全时，缺的部分回退到 DEFAULT_FORMAT，行为与改造前完全一致（零回归）。
DEFAULT_FORMAT = {
    "页面": {"纸张宽cm": 21.0, "纸张高cm": 29.7, "上边距cm": 2.5, "下边距cm": 2.0,
             "左边距cm": 2.0, "右边距cm": 2.0, "页眉距cm": 1.5, "页脚距cm": 1.75},
    "字体": {"中文字体": "宋体", "西文字体": "Times New Roman", "字体颜色": "000000"},
    "正文": {"字号pt": 12.0, "行距倍数": 1.5, "首行缩进字符": 2, "对齐": "两端对齐"},
    "标题": {"字号pt": 16.0, "加粗": True, "对齐": "居中", "段前行": 1.0, "段后行": 1.0,
             "行距倍数": 1.5, "编号格式": "decimal", "一级标题另起页": True},
    "序号标题": {"字号pt": 12.0, "加粗": False, "行距倍数": 1.5, "首行缩进字符": 2, "左缩进字符": 0},
    "列表": {"字号pt": 12.0, "加粗": False, "行距倍数": 1.5, "首行缩进字符": 2, "每级左缩进字符": 2},
    "表格": {"表题前缀": "表", "表题字号pt": 10.5, "表题加粗": False, "表题段前行": 0.5,
             "内容字号pt": 10.5, "表头加粗": True, "边框粗细": 4, "宽度百分比": 100},
    "配图": {"图题前缀": "图", "图题字号pt": 10.5, "图题加粗": True, "图题段后行": 0.5,
             "图前段前行": 0.5, "默认宽度cm": 12.0, "最大宽度cm": 17.0,
             "占位文本": "【此处为占位图，请替换为实际截图】"},
    "目录": {"标题文本": "目 录", "标题字号pt": 16.0, "级别范围": "1-3"},
    "页码": {"字号pt": 10.5, "对齐": "居中", "起始页码": 1},
}

_CFG = copy.deepcopy(DEFAULT_FORMAT)

_ALIGN = {"居中": WD_ALIGN_PARAGRAPH.CENTER, "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
          "左对齐": WD_ALIGN_PARAGRAPH.LEFT, "右对齐": WD_ALIGN_PARAGRAPH.RIGHT}


def _deep_merge(base, over):
    for k, v in (over or {}).items():
        if k == "_说明":
            continue
        if isinstance(v, dict) and isinstance(base.get(k), dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v
    return base


def cfg(group, key):
    """读一个排版参数；JSON 缺该键时回退到 DEFAULT_FORMAT。"""
    g = _CFG.get(group, {})
    if key in g:
        return g[key]
    return DEFAULT_FORMAT.get(group, {}).get(key)


def align(group, key="对齐"):
    return _ALIGN.get(cfg(group, key), WD_ALIGN_PARAGRAPH.JUSTIFY)


def load_format_config(path):
    """加载排版配置 JSON（深合并到内置默认）并刷新派生的字体/字号常量。"""
    global _CFG, SZ_SAN, SZ_XIAOSI, SZ_WU, CN_FONT, EN_FONT, BLACK
    _CFG = copy.deepcopy(DEFAULT_FORMAT)
    if path and os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                _deep_merge(_CFG, json.load(f))
            print(f"[排版配置] 已加载：{path}")
        except (json.JSONDecodeError, OSError) as e:
            print(f"[排版配置] 读取 {path} 失败（{e}），改用内置默认排版。")
    else:
        print("[排版配置] 未找到 format_config.json，使用内置默认排版。")
    CN_FONT = cfg("字体", "中文字体")
    EN_FONT = cfg("字体", "西文字体")
    try:
        BLACK = RGBColor.from_string(str(cfg("字体", "字体颜色")).lstrip("#"))
    except ValueError:
        BLACK = RGBColor(0, 0, 0)
    SZ_SAN = Pt(cfg("标题", "字号pt"))       # 三号（标题）
    SZ_XIAOSI = Pt(cfg("正文", "字号pt"))     # 小四（正文）
    SZ_WU = Pt(cfg("表格", "内容字号pt"))     # 五号（表格内容，题注/页码另读各自配置）
    return _CFG


def _discover_config(content_path, cli_path):
    """定位排版配置：命令行指定 > content.json 同目录 > 本脚本同目录。"""
    if cli_path:
        return cli_path
    for d in (os.path.dirname(os.path.abspath(content_path)),
              os.path.dirname(os.path.abspath(__file__))):
        cand = os.path.join(d, "format_config.json")
        if os.path.exists(cand):
            return cand
    return None


# 派生字体/字号常量（load_format_config 会按配置刷新；此处给内置默认，保证未加载时也可用）
SZ_SAN = Pt(16)      # 三号
SZ_XIAOSI = Pt(12)   # 小四
SZ_WU = Pt(10.5)     # 五号
CN_FONT = "宋体"
EN_FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)   # 全文字体统一为纯黑


# ============================ 底层 XML 工具 ============================
def set_run_fonts(run, cn=None, en=None):
    """设置一个 run 的中文字体(eastAsia)与西文字体(ascii/hAnsi)。
    默认取当前全局 CN_FONT/EN_FONT（load_format_config 后已按配置刷新），
    不用默认参数绑定，避免加载配置前后取到旧字体。"""
    if cn is None:
        cn = CN_FONT
    if en is None:
        en = EN_FONT
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:cs'), en)
    run.font.color.rgb = BLACK   # 强制纯黑
    run.italic = False           # 强制非斜体
    run.font.italic = False


def set_first_line_indent_chars(paragraph, chars=2):
    """首行缩进 N 字符 (firstLineChars)。chars=0 取消缩进。
    通过 paragraph_format 创建 w:ind，保证 pPr 子元素顺序合规，再补字符属性。"""
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(chars * 12) if chars and chars > 0 else Pt(0)
    ind = paragraph._p.get_or_add_pPr().find(qn('w:ind'))
    if ind is not None:
        if chars and chars > 0:
            ind.set(qn('w:firstLineChars'), str(int(chars * 100)))
        else:
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')


def set_left_indent_chars(paragraph, chars):
    """左缩进 N 字符 (用于列表层级)。"""
    pf = paragraph.paragraph_format
    pf.left_indent = Pt(chars * 12)
    ind = paragraph._p.get_or_add_pPr().find(qn('w:ind'))
    if ind is not None:
        ind.set(qn('w:leftChars'), str(int(chars * 100)))


def set_spacing_lines(paragraph, before_lines=0.0, after_lines=0.0, line_spacing=1.5):
    """段前/段后(单位:行) + 行距倍数。用 paragraph_format 创建 w:spacing(顺序合规)，
    再补 beforeLines/afterLines(单位:百分之一行) 让 Word 以"行"为单位显示间距。"""
    pf = paragraph.paragraph_format
    pf.line_spacing = float(line_spacing)              # 浮点=倍数行距 (lineRule=auto)
    pf.space_before = Pt(before_lines * 12)
    pf.space_after = Pt(after_lines * 12)
    spc = paragraph._p.get_or_add_pPr().find(qn('w:spacing'))
    if spc is not None:
        spc.set(qn('w:beforeLines'), str(int(before_lines * 100)))
        spc.set(qn('w:afterLines'), str(int(after_lines * 100)))


def add_field(paragraph, instr, default_text=""):
    """插入一个 Word 域 (begin/instrText/separate/end)。用于 TOC 和 PAGE。"""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldSep)

    run4 = paragraph.add_run(default_text)

    run5 = paragraph.add_run()
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldEnd)
    return [run, run2, run3, run4, run5]


def set_cell_borders_all(cell):
    """单元格四边全框线 (single 黑，粗细取配置『表格.边框粗细』，单位1/8磅)。"""
    sz = str(int(cfg("表格", "边框粗细")))
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)


def set_table_autofit_window(table):
    """表格【根据窗口调整】：宽度=版心×配置百分比，禁止超出页边距。"""
    pct = str(int(float(cfg("表格", "宽度百分比")) * 50))   # pct 单位为1/50%，100%→5000
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), pct); tblW.set(qn('w:type'), 'pct')
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'autofit')
    tblPr.append(layout)


def update_fields_on_open(document):
    """在 settings.xml 写入 updateFields，使 Word 打开时提示更新目录/页码域。"""
    settings = document.settings.element
    uf = settings.find(qn('w:updateFields'))
    if uf is None:
        uf = OxmlElement('w:updateFields')
        settings.append(uf)
    uf.set(qn('w:val'), 'true')


# ============================ 样式与节设置 ============================
def setup_base_styles(document):
    """正文 Normal 样式: 宋体/Times New Roman 小四 1.5倍行距 两端对齐 首行缩进2字符。"""
    normal = document.styles['Normal']
    normal.font.name = EN_FONT
    normal.font.size = SZ_XIAOSI
    normal.font.color.rgb = BLACK
    normal.font.italic = False
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), EN_FONT)
    rfonts.set(qn('w:hAnsi'), EN_FONT)
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    rfonts.set(qn('w:cs'), EN_FONT)


def setup_heading_styles(document):
    """配置 Heading 1~5 样式匹配排版规范：宋体三号加粗居中，段前段后1行，行距1.5倍。
    创建 Heading 4/5（默认模板可能不含），使样式面板可联动识别。"""
    for level in range(1, 6):
        style_name = f"Heading {level}"
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = EN_FONT
        style.font.size = SZ_SAN
        style.font.bold = bool(cfg("标题", "加粗"))
        style.font.color.rgb = BLACK
        style.font.italic = False
        pf = style.paragraph_format
        pf.alignment = align("标题")
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(float(cfg("标题", "段前行")) * 12)
        pf.space_after = Pt(float(cfg("标题", "段后行")) * 12)
        pf.line_spacing = float(cfg("标题", "行距倍数"))
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:ascii'), EN_FONT)
        rfonts.set(qn('w:hAnsi'), EN_FONT)
        rfonts.set(qn('w:eastAsia'), CN_FONT)
        rfonts.set(qn('w:cs'), EN_FONT)
        pf.outline_level = level - 1


def setup_heading_numbering(document):
    """创建混合多级列表编号定义（hybridMultilevel）用于标题自动编号。
    格式: 1 / 1.1 / 1.1.1 / 1.1.1.1 / 1.1.1.1.1
    返回 numId，render_heading 需将此 id 挂到段落上。"""
    numbering = document.part.numbering_part._element
    existing_abs = [int(el.get(qn('w:abstractNumId'))) for el in numbering.findall(qn('w:abstractNum'))]
    existing_num = [int(el.get(qn('w:numId'))) for el in numbering.findall(qn('w:num'))]
    new_abs_id = max(existing_abs) + 1 if existing_abs else 0
    new_num_id = max(existing_num) + 1 if existing_num else 1
    abs_num = OxmlElement('w:abstractNum')
    abs_num.set(qn('w:abstractNumId'), str(new_abs_id))
    mlt = OxmlElement('w:multiLevelType')
    mlt.set(qn('w:val'), 'hybridMultilevel')
    abs_num.append(mlt)
    for lvl in range(5):
        lvl_elem = OxmlElement('w:lvl')
        lvl_elem.set(qn('w:ilvl'), str(lvl))
        for tag, val in [('start', '1'), ('numFmt', str(cfg("标题", "编号格式")))]:
            el = OxmlElement(f'w:{tag}')
            el.set(qn('w:val'), val)
            lvl_elem.append(el)
        lvl_txt = OxmlElement('w:lvlText')
        lvl_txt.set(qn('w:val'), '.'.join(f'%{i+1}' for i in range(lvl + 1)))
        lvl_elem.append(lvl_txt)
        jc = OxmlElement('w:lvlJc')
        jc.set(qn('w:val'), 'left')
        lvl_elem.append(jc)
        rPr = OxmlElement('w:rPr')
        rfonts = OxmlElement('w:rFonts')
        for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
            rfonts.set(qn(attr), EN_FONT if attr != 'w:eastAsia' else CN_FONT)
        rPr.append(rfonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(round(float(cfg("标题", "字号pt")) * 2))))  # 半磅
        rPr.append(sz)
        if cfg("标题", "加粗"):
            rPr.append(OxmlElement('w:b'))
        lvl_elem.append(rPr)
        lvl_elem.append(OxmlElement('w:suff'))
        lvl_elem.find(qn('w:suff')).set(qn('w:val'), 'space')
        abs_num.append(lvl_elem)
    numbering.append(abs_num)
    num_elem = OxmlElement('w:num')
    num_elem.set(qn('w:numId'), str(new_num_id))
    ref = OxmlElement('w:abstractNumId')
    ref.set(qn('w:val'), str(new_abs_id))
    num_elem.append(ref)
    numbering.append(num_elem)
    return new_num_id


def set_page_layout(section):
    """纸张与页边距，取配置『页面』组（默认 A4 + 上2.5cm其余2cm）。"""
    section.page_width = Cm(float(cfg("页面", "纸张宽cm")))
    section.page_height = Cm(float(cfg("页面", "纸张高cm")))
    section.top_margin = Cm(float(cfg("页面", "上边距cm")))
    section.bottom_margin = Cm(float(cfg("页面", "下边距cm")))
    section.left_margin = Cm(float(cfg("页面", "左边距cm")))
    section.right_margin = Cm(float(cfg("页面", "右边距cm")))
    section.header_distance = Cm(float(cfg("页面", "页眉距cm")))
    section.footer_distance = Cm(float(cfg("页面", "页脚距cm")))


def restart_page_numbering(section, start=1):
    """该节页码重新从 start 开始 (阿拉伯数字)。"""
    sectPr = section._sectPr
    pgNum = sectPr.find(qn('w:pgNumType'))
    if pgNum is None:
        pgNum = OxmlElement('w:pgNumType')
        cols = sectPr.find(qn('w:cols'))   # pgNumType 必须排在 cols 之前
        if cols is not None:
            cols.addprevious(pgNum)
        else:
            sectPr.append(pgNum)
    pgNum.set(qn('w:fmt'), 'decimal')
    pgNum.set(qn('w:start'), str(start))


def clear_footer(section):
    """清空页脚 (封面/目录不编号)。"""
    section.footer.is_linked_to_previous = False
    for p in section.footer.paragraphs:
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        p.text = ""


def add_centered_page_number(section):
    """页脚页码（对齐/字号取配置『页码』组，默认居中五号 Times New Roman）。"""
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    p.alignment = align("页码")
    set_spacing_lines(p, 0, 0, 1.0)
    set_first_line_indent_chars(p, 0)
    runs = add_field(p, ' PAGE ', '1')
    for r in runs:
        r.font.size = Pt(cfg("页码", "字号pt"))
        set_run_fonts(r, cn=EN_FONT, en=EN_FONT)


# ============================ 内容块渲染 ============================
CIRCLED = [chr(0x2460 + i) for i in range(20)]  # ①..⑳


def render_heading(document, level, text, num_id):
    """五级标题: 宋体三号加粗居中，段前段后1行
    使用 Word 内置标题样式（Heading 1~5）+ 多级列表自动编号，样式与样式面板联动。
    一级标题的另起一页由调用方在返回的 paragraph 上设 page_break_before。"""
    style_name = f"Heading {level}"
    p = document.add_paragraph()
    p.style = document.styles[style_name]
    p.alignment = align("标题")
    set_first_line_indent_chars(p, 0)
    set_spacing_lines(p, before_lines=float(cfg("标题", "段前行")),
                      after_lines=float(cfg("标题", "段后行")),
                      line_spacing=float(cfg("标题", "行距倍数")))
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), str(level - 1))
    pPr.append(outline)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level - 1))
    numPr.append(ilvl)
    numIdElem = OxmlElement('w:numId')
    numIdElem.set(qn('w:val'), str(num_id))
    numPr.append(numIdElem)
    pPr.append(numPr)
    run = p.add_run(text)
    run.bold = bool(cfg("标题", "加粗"))
    run.font.size = Pt(cfg("标题", "字号pt"))
    set_run_fonts(run)
    return p


def render_paragraph(document, text):
    """正文段落。"""
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = align("正文")
    set_first_line_indent_chars(p, int(cfg("正文", "首行缩进字符")))
    set_spacing_lines(p, 0, 0, float(cfg("正文", "行距倍数")))
    run = p.add_run(text)
    run.font.size = Pt(cfg("正文", "字号pt"))
    set_run_fonts(run)
    return p


def render_hnum(document, seq, text, level=0):
    """内联序号标题（末级标题降级 / 叶内小标题）：正文样式，不加粗、小四宋体、
    1.5倍行距，不参与标题自动编号、不设大纲级别、不进目录。排版与正文一致。
    按 marker 层选序号样式：level0→（n）、level1→n）、level≥2→①，并按层左缩进。"""
    if level <= 0:
        marker = f"（{seq}）"
    elif level == 1:
        marker = f"{seq}）"
    else:
        marker = CIRCLED[min(seq - 1, 19)]
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = align("正文")
    # 序号标题的缩进与正文完全一致：左侧0字符、首行2字符（不按 marker 层做左缩进，
    # 否则 1）/①这类标题会比正文多缩进一截，与正文错位）。
    set_first_line_indent_chars(p, int(cfg("序号标题", "首行缩进字符")))
    set_left_indent_chars(p, int(cfg("序号标题", "左缩进字符")))
    set_spacing_lines(p, 0, 0, float(cfg("序号标题", "行距倍数")))
    run = p.add_run(f"{marker}{text}")
    run.bold = bool(cfg("序号标题", "加粗"))
    run.font.size = Pt(cfg("序号标题", "字号pt"))
    set_run_fonts(run)
    return p


def render_list(document, items, level=0, counters=None):
    """三级列表: （1） / 1） / ① ；小四宋体不加粗，首行缩进，1.5倍行距。"""
    if counters is None:
        counters = [0, 0, 0]
    counters[level] = 0
    for d in range(level + 1, 3):
        counters[d] = 0
    for item in items:
        counters[level] += 1
        n = counters[level]
        if level == 0:
            marker = f"（{n}）"
        elif level == 1:
            marker = f"{n}）"
        else:
            marker = CIRCLED[min(n - 1, 19)]
        p = document.add_paragraph()
        p.style = document.styles['Normal']
        p.alignment = align("正文")
        set_first_line_indent_chars(p, int(cfg("列表", "首行缩进字符")))
        set_left_indent_chars(p, level * int(cfg("列表", "每级左缩进字符")))
        set_spacing_lines(p, 0, 0, float(cfg("列表", "行距倍数")))
        text = item.get('t', '') if isinstance(item, dict) else str(item)
        has_children = isinstance(item, dict) and item.get('children') and level < 2
        text = text.rstrip()
        if has_children:
            # 列表标题（其下还有子列表）：去掉末尾句号/分号
            text = text.rstrip("。；;")
        elif text and text[-1] not in "。！？":
            # 叶子罗列项：确保以句号收尾（先去掉尾部逗顿分号再补句号）
            text = text.rstrip("；;，,、 ") + "。"
        run = p.add_run(marker + text)
        run.bold = bool(cfg("列表", "加粗"))
        run.font.size = Pt(cfg("列表", "字号pt"))
        set_run_fonts(run)
        if has_children:
            render_list(document, item['children'], level + 1, counters)


def render_table(document, chapter, tbl_seq, title, header, rows):
    """全框线表格 + 表题在上(五号宋体)。表头加粗，内容居中，五号字。分章编号。"""
    cap = document.add_paragraph()
    cap.style = document.styles['Normal']
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(cap, 0)
    set_spacing_lines(cap, before_lines=float(cfg("表格", "表题段前行")),
                      after_lines=0.0, line_spacing=1.5)
    crun = cap.add_run(f"{cfg('表格', '表题前缀')}{chapter}-{tbl_seq} {title}")
    crun.bold = bool(cfg("表格", "表题加粗"))
    crun.font.size = Pt(cfg("表格", "表题字号pt"))
    set_run_fonts(crun)

    ncols = len(header)
    table = document.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_autofit_window(table)

    def fill_cell(cell, text, bold=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders_all(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_first_line_indent_chars(para, 0)
        set_spacing_lines(para, 0, 0, 1.0)
        run = para.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(cfg("表格", "内容字号pt"))
        set_run_fonts(run)

    for j, h in enumerate(header):
        fill_cell(table.rows[0].cells[j], h, bold=bool(cfg("表格", "表头加粗")))
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            fill_cell(table.rows[i + 1].cells[j], val, bold=False)
    return table


def render_figure(document, chapter, fig_seq, title, img_path=None, width_cm=None):
    """图片居中无缩进 + 图题在下(五号宋体加粗)。分章编号。无图则放占位框。"""
    if width_cm is None:
        width_cm = float(cfg("配图", "默认宽度cm"))
    p = document.add_paragraph()
    p.style = document.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(p, 0)
    set_spacing_lines(p, before_lines=float(cfg("配图", "图前段前行")),
                      after_lines=0.0, line_spacing=1.5)
    run = p.add_run()
    if img_path and os.path.exists(img_path):
        run.add_picture(img_path, width=Cm(min(width_cm, float(cfg("配图", "最大宽度cm")))))
    else:
        ph = run.add_text(str(cfg("配图", "占位文本")))
        run.font.size = SZ_WU
        set_run_fonts(run)
    cap = document.add_paragraph()
    cap.style = document.styles['Normal']
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(cap, 0)
    set_spacing_lines(cap, before_lines=0.0,
                      after_lines=float(cfg("配图", "图题段后行")), line_spacing=1.5)
    crun = cap.add_run(f"{cfg('配图', '图题前缀')}{chapter}-{fig_seq} {title}")
    crun.bold = bool(cfg("配图", "图题加粗"))
    crun.font.size = Pt(cfg("配图", "图题字号pt"))
    set_run_fonts(crun)


# ============================ 封面 / 目录 ============================
def build_cover(document, meta):
    """简洁封面: 项目名 / 投标人 / 日期，全部居中。"""
    for _ in range(4):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(title, 0)
    set_spacing_lines(title, 1, 1, 1.5)
    r = title.add_run(meta.get('project_name', '投标文件'))
    r.bold = True; r.font.size = Pt(26)
    set_run_fonts(r)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(sub, 0)
    set_spacing_lines(sub, 1, 1, 1.5)
    r = sub.add_run(meta.get('doc_type', '技术投标文件'))
    r.bold = True; r.font.size = Pt(22)
    set_run_fonts(r)

    for _ in range(6):
        document.add_paragraph()

    for label, key in [("投标人：", "bidder"), ("招标编号：", "tender_no"), ("日期：", "date")]:
        if meta.get(key):
            line = document.add_paragraph()
            line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_first_line_indent_chars(line, 0)
            set_spacing_lines(line, 0.5, 0.5, 1.5)
            r = line.add_run(f"{label}{meta.get(key)}")
            r.font.size = SZ_XIAOSI
            set_run_fonts(r)


def build_toc(document):
    """目录页: '目 录' 三号宋体加粗居中 + 自动目录域(1-3级)。"""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent_chars(title, 0)
    set_spacing_lines(title, 0.5, 1.0, 1.5)
    r = title.add_run(str(cfg("目录", "标题文本")))
    r.bold = True; r.font.size = Pt(cfg("目录", "标题字号pt"))
    set_run_fonts(r)

    p = document.add_paragraph()
    set_first_line_indent_chars(p, 0)
    set_spacing_lines(p, 0, 0, 1.5)
    lvl_range = str(cfg("目录", "级别范围"))
    runs = add_field(p, f' TOC \\o "{lvl_range}" \\h \\z \\u ', "右键此处选择“更新域”以生成目录")
    for rr in runs:
        rr.font.size = SZ_XIAOSI
        set_run_fonts(rr)


# ============================ 校验接入 ============================
def _run_validation(data, base_dir):
    """渲染前调用 validate_content 校验。返回 True 表示可继续，False 表示有错误应中止。"""
    try:
        from validate_content import validate
    except ImportError:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        try:
            from validate_content import validate
        except ImportError:
            print("提示: 未找到 validate_content.py，跳过校验。")
            return True
    errors, warns = validate(data, base_dir)
    for w in warns:
        print(f"[警告] {w}")
    for e in errors:
        print(f"[错误] {e}")
    print(f"校验: {len(errors)} 个错误，{len(warns)} 个警告。")
    if errors:
        print("已中止：请修正上述错误后重试，或加 --no-validate 强行生成 (不推荐)。")
        return False
    return True


# ============================ 主流程 ============================
def _swap_cjk_quotes(s):
    """直角引号统一改为中文引号：「」→“”、『』→‘’。《》书名号保留不动。"""
    return (s.replace("「", "“").replace("」", "”")
             .replace("『", "‘").replace("』", "’"))


def _deep_normalize_quotes(obj):
    """递归把 content 里所有字符串的直角引号规范化——渲染层兜底，
    不论 content.json 新旧，成品 docx 一律用中文引号。"""
    if isinstance(obj, str):
        return _swap_cjk_quotes(obj)
    if isinstance(obj, list):
        return [_deep_normalize_quotes(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _deep_normalize_quotes(v) for k, v in obj.items()}
    return obj


def build(content_path, output_path, do_validate=True, format_config=None):
    load_format_config(_discover_config(content_path, format_config))
    with open(content_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    data = _deep_normalize_quotes(data)   # 直角引号→中文引号（渲染层兜底）

    if do_validate and not _run_validation(data, os.path.dirname(os.path.abspath(content_path))):
        sys.exit(1)

    meta = data.get('meta', {})
    body = data.get('body', [])
    base_dir = os.path.dirname(os.path.abspath(content_path))

    document = Document()
    setup_base_styles(document)
    setup_heading_styles(document)
    heading_num_id = setup_heading_numbering(document)

    # —— 第1节: 目录（一律不做封面）——
    sec0 = document.sections[0]
    set_page_layout(sec0)
    build_toc(document)

    # —— 第2节: 正文 (页码从1重启) ——
    document.add_section(WD_SECTION.NEW_PAGE)
    sec1 = document.sections[-1]
    set_page_layout(sec1)
    restart_page_numbering(sec1, start=int(cfg("页码", "起始页码")))

    clear_footer(sec0)
    add_centered_page_number(sec1)

    h1_counter = 0
    tbl_counter = {}
    fig_counter = {}
    first_h1 = True

    for block in body:
        btype = block.get('type')
        if btype in ('h1', 'h2', 'h3', 'h4', 'h5'):
            level = int(btype[1])
            if level == 1:
                h1_counter += 1
            p = render_heading(document, level, block.get('text', ''), heading_num_id)
            if level == 1:
                if not first_h1 and cfg("标题", "一级标题另起页"):
                    p.paragraph_format.page_break_before = True
                first_h1 = False
        elif btype == 'hnum':
            render_hnum(document, block.get('seq', 1), block.get('text', ''),
                        block.get('level', 0))
        elif btype == 'p':
            render_paragraph(document, block.get('text', ''))
        elif btype == 'list':
            render_list(document, block.get('items', []))
        elif btype == 'table':
            ch = h1_counter if h1_counter > 0 else 1
            tbl_counter[ch] = tbl_counter.get(ch, 0) + 1
            render_table(document, ch, tbl_counter[ch],
                         block.get('title', ''), block.get('header', []),
                         block.get('rows', []))
        elif btype == 'figure':
            ch = h1_counter if h1_counter > 0 else 1
            fig_counter[ch] = fig_counter.get(ch, 0) + 1
            img = block.get('img')
            if img and not os.path.isabs(img):
                img = os.path.join(base_dir, img)
            w = block.get('width_cm')
            render_figure(document, ch, fig_counter[ch],
                          block.get('title', ''), img,
                          float(w) if w is not None else None)

    update_fields_on_open(document)
    document.save(output_path)
    print(f"已生成: {output_path}")
    print(f"标题数: {sum(1 for b in body if b.get('type','') in ('h1','h2','h3','h4','h5'))} | "
          f"表: {sum(v for v in tbl_counter.values())} | "
          f"图: {sum(v for v in fig_counter.values())}")
    print("提示: 用 Word 打开后按 Ctrl+A 再按 F9 更新目录与页码域。")


if __name__ == '__main__':
    raw = sys.argv[1:]
    do_validate = '--no-validate' not in raw
    format_config = None
    args = []
    i = 0
    while i < len(raw):
        a = raw[i]
        if a == '--no-validate':
            i += 1
        elif a == '--format-config':
            format_config = raw[i + 1] if i + 1 < len(raw) else None
            i += 2
        elif a.startswith('--format-config='):
            format_config = a.split('=', 1)[1]
            i += 1
        else:
            args.append(a)
            i += 1
    if len(args) != 2:
        print("用法: python build_docx.py content.json output.docx "
              "[--no-validate] [--format-config 路径]")
        sys.exit(1)
    build(args[0], args[1], do_validate=do_validate, format_config=format_config)
