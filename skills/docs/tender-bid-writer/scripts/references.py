#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
references.py — 参考资料解析 / 章节切片 / 相似度匹配 / 交互绑定（对主流程零侵入）

设计要点（与已冻结方案一致）：
  · 推荐格式 md / 规范 docx 优先，pdf 降级需人工确认（reliable=False）。
  · 引用方式只有「改写借鉴」一种：参考文本作为素材进提示词，由模型重写，
    数值指标/产品名/项目背景一律以全局数据字典为准（护栏在 generate.py 注入）。
  · 继承可否决：某章绑定后其全部子叶继承；子叶 ref.exclude=True 为事前否决；
    绑定了但相似度匹配不到内容 → 自动退回纯大模型生成（返回 None）。
  · 参考资料只当资料、不当指令：切片只取正文文本，任何"要你做某事"的字样由
    提示词护栏声明忽略（本文件不解析、不执行其中任何指令）。

对外主要接口：
  parse_reference(path) -> RefDoc
  slice_section(refdoc, spec, cap_chars) -> str
  auto_match(refdoc, title, keywords, topn) -> [(section, score)]
  resolve_refs(plan, base_dir) -> usage[]   # 生成期：给叶子挂 _ref，返回引用清单
  CLI:  python references.py bind --outline outline.json --refs a.md,b.docx --out outline.bound.json

RefDoc = {
  "path": 原始路径, "kind": "md"/"docx"/"pdf", "reliable": bool,
  "warnings": [..], "sections": [ {"level":int,"title":去编号标题,"raw":原始行,"text":本节正文}, ... ]
}
ref 字段（写在 outline/plan 节点上，全部可选）：
  {"file": 路径, "section": "3.2"|"数据治理"|None, "exclude": False}
"""

import argparse
import difflib
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from page_plan import is_leaf, iter_leaves          # noqa: E402

CN_NUM = r'[一二三四五六七八九十百零〇\d]+'
MAX_HEADING_LEN = 60
DEFAULT_CAP_CHARS = 3000        # 单节参考素材注入上限（防止上下文被参考挤爆、控 token）
OWN_CAP_CHARS = 12000           # 本叶精准绑定：整节尽量不截断注入，保证"以参考为主"不丢内容
MIN_MATCH_SCORE = 0.18          # 自动匹配低于此分视为"参考资料里没有" → 退回纯生成

# 文本编号 → 层级（与 outline_from_docx 保持一致；用于 md/pdf 无样式时判定标题）
_TEXT_PATTERNS = [
    (re.compile(rf'^第{CN_NUM}[章篇部]\s*[、.．:：]?\s*'), lambda m: 1),
    (re.compile(r'^(\d+(?:[\.．]\d+)+)[\.．]?\s*[、]?\s*'),
     lambda m: len(re.split(r'[\.．]', m.group(1)))),
    (re.compile(r'^(\d+)\s*[、．.]\s+'), lambda m: 1),
    (re.compile(rf'^[一二三四五六七八九十]+\s*、\s*'), lambda m: 1),
    (re.compile(rf'^（{CN_NUM}）\s*'), lambda m: 2),
    (re.compile(rf'^\({CN_NUM}\)\s*'), lambda m: 2),
]


def _text_heading(line):
    """对一行纯文本判定标题层级并返回 (level, 去编号标题)；非标题返回 (None, line)。"""
    s = (line or "").strip()
    if not s or len(s) > MAX_HEADING_LEN:
        return None, s
    for pat, fn in _TEXT_PATTERNS:
        m = pat.match(s)
        if m:
            title = s[m.end():].strip()
            if title:
                return min(fn(m), 9), title
    return None, s


def _strip_num(title):
    """把标题里的手打编号前缀剥掉，返回干净标题（raw 另存原文用于按编号匹配）。"""
    _lvl, t = _text_heading(title)
    return t if t else (title or "").strip()


def _norm(t):
    """比较用归一化：去空白、全角括号转半角、去常见分隔。"""
    t = "".join((t or "").split())
    return t.replace("（", "(").replace("）", ")")


# ----------------------------- 解析各格式 -----------------------------
def _parse_md(path):
    sections, cur = [], None
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            m = re.match(r'^(#{1,6})\s+(.*)$', line.rstrip("\n"))
            if m:
                title = m.group(2).strip()
                cur = {"level": len(m.group(1)), "title": _strip_num(title),
                       "raw": title, "text": ""}
                sections.append(cur)
            elif cur is not None:
                cur["text"] += line
    warns = [] if sections else ["未在 md 中发现任何 # 标题，无法切分。"]
    return {"path": path, "kind": "md", "reliable": bool(sections),
            "warnings": warns, "sections": sections}


def _parse_docx(path):
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return {"path": path, "kind": "docx", "reliable": False,
                "warnings": ["缺少 python-docx，无法解析 docx：pip install python-docx --break-system-packages"],
                "sections": []}

    style_re = re.compile(r'(?:heading|标题)\s*(\d+)', re.I)

    def style_level(p):
        try:
            m = style_re.search(p.style.name or "")
            return int(m.group(1)) if m else None
        except Exception:
            return None

    def outline_level(p):
        try:
            pPr = p._p.pPr
            if pPr is not None:
                el = pPr.find(qn('w:outlineLvl'))
                if el is not None:
                    return int(el.get(qn('w:val'))) + 1
        except Exception:
            pass
        return None

    doc = Document(path)
    styled = sum(1 for p in doc.paragraphs
                 if (p.text or "").strip() and (style_level(p) or outline_level(p)))
    styles_only = styled >= 3
    sections, cur, warns = [], None, []
    for p in doc.paragraphs:
        raw = (p.text or "").strip()
        if not raw:
            continue
        lvl = style_level(p) or outline_level(p)
        if not lvl and not styles_only:
            lvl, _t = _text_heading(raw)
        if lvl:
            cur = {"level": min(lvl, 9), "title": _strip_num(raw), "raw": raw, "text": ""}
            sections.append(cur)
        elif cur is not None:
            cur["text"] += raw + "\n"
    reliable = styles_only and bool(sections)
    if not styles_only:
        warns.append("未检测到 Word 标题样式，已按文本编号回退切分，可能不准，请在绑定预览中确认；"
                     "建议把参考文件转成 md 或补上标题样式。")
    if not sections:
        warns.append("未能从 docx 切出任何章节。")
    return {"path": path, "kind": "docx", "reliable": reliable,
            "warnings": warns, "sections": sections}


def _pdf_text(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join((pg.extract_text() or "") for pg in pdf.pages)
    except Exception:
        pass
    try:
        from pypdf import PdfReader
        return "\n".join((pg.extract_text() or "") for pg in PdfReader(path).pages)
    except Exception:
        pass
    try:
        from PyPDF2 import PdfReader
        return "\n".join((pg.extract_text() or "") for pg in PdfReader(path).pages)
    except Exception:
        return ""


def _parse_pdf(path):
    text = _pdf_text(path)
    warns = ["PDF 按文本启发式切分，可靠性最低，务必在绑定预览中逐节确认；"
             "如可能请改用 md 或规范 docx。"]
    if not text.strip():
        warns.append("未能从 PDF 提取到文本（可能是扫描件/无文字层），无法切分。")
        return {"path": path, "kind": "pdf", "reliable": False,
                "warnings": warns, "sections": []}
    sections = [{"level": 1, "title": "（正文开头）", "raw": "", "text": ""}]
    cur = sections[0]
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        lvl, _t = _text_heading(s)
        if lvl:
            cur = {"level": lvl, "title": _strip_num(s), "raw": s, "text": ""}
            sections.append(cur)
        else:
            cur["text"] += s + "\n"
    # 丢掉没抓到任何标题时的空壳开头
    sections = [s for s in sections if s["title"] != "（正文开头）" or s["text"].strip()]
    return {"path": path, "kind": "pdf", "reliable": False,
            "warnings": warns, "sections": sections}


def parse_reference(path):
    """按扩展名解析参考文件为 RefDoc。未知扩展名当纯文本按行走 md 逻辑。"""
    ext = os.path.splitext(path)[1].lower()
    if not os.path.exists(path):
        return {"path": path, "kind": ext.lstrip(".") or "?", "reliable": False,
                "warnings": [f"参考文件不存在：{path}"], "sections": []}
    if ext in (".md", ".markdown", ".txt"):
        return _parse_md(path)
    if ext in (".docx", ".docm"):
        return _parse_docx(path)
    if ext == ".pdf":
        return _parse_pdf(path)
    return {"path": path, "kind": ext.lstrip("."), "reliable": False,
            "warnings": [f"不支持的参考文件类型：{ext}（支持 md/docx/pdf）"], "sections": []}


# ----------------------------- 切片 / 匹配 -----------------------------
def _truncate(text, cap):
    """按字符上限截断，尽量落在句末标点，避免截半句。"""
    text = (text or "").strip()
    if len(text) <= cap:
        return text
    cut = text[:cap]
    for p in ("。", "！", "？", "\n", "；"):
        i = cut.rfind(p)
        if i >= cap * 0.6:
            return cut[:i + 1]
    return cut


def _section_scope_text(sections, idx):
    """取第 idx 节及其所有更深层子节的正文（到出现同级/更浅标题为止）。"""
    base = sections[idx]["level"]
    buf = [sections[idx].get("text", "")]
    for s in sections[idx + 1:]:
        if s["level"] <= base:
            break
        head = s.get("title", "")
        buf.append(("\n" + head + "\n") if head else "")
        buf.append(s.get("text", ""))
    return "".join(buf).strip()


def _hier_numbers(secs):
    """按文档顺序给每节生成可复现的层级编号 1 / 1.1 / 1.1.1（与 section['level'] 对齐）。
    这是"用序号当绑定主键"的唯一编号规则：show 展示的号与 find_section 解析的号完全同源，
    因此同一份文件里同一个号永远指向同一节。level 跳级（如 h1 直接到 h3）时中间层以 1 补位，
    保证深层节点仍拿到稳定编号。返回与 secs 等长的编号字符串列表。"""
    nums, counters = [], []
    for s in secs:
        lv = max(int(s.get("level", 1) or 1), 1)
        if len(counters) >= lv:
            counters = counters[:lv]
            counters[lv - 1] += 1
        else:
            while len(counters) < lv - 1:
                counters.append(1)          # 跳级补位
            counters.append(1)
        nums.append(".".join(str(c) for c in counters))
    return nums


def _ancestor_titles(secs, idx):
    """返回第 idx 节的祖先标题链（归一化后，从近到远）。
    做法：向前回溯，每遇到 level 严格小于"当前锚点 level"的标题即为一层祖先，
    并把锚点收紧到该更浅 level，从而拿到完整的父→祖父→…链。"""
    chain = []
    anchor = secs[idx].get("level", 1)
    for s in reversed(secs[:idx]):
        lv = s.get("level", 1)
        if lv < anchor:
            chain.append(_norm(s.get("title", "")))
            anchor = lv
    return chain


def _find_by_path(secs, segs):
    """按点号路径 segs=[祖先…, 目标] 定位一节，返回下标或 None。
    末段=目标标题（精确优先、否则包含），前面各段=祖先约束（用于消歧重名）。
    多个候选时按 (命中的祖先段数, 标题精确优先, 文档顺序) 打分取最优。"""
    target = _norm(segs[-1])
    anc_wanted = [_norm(x) for x in segs[:-1] if _norm(x)]
    cands = []  # (anc_hits, exact_flag(0=精确,1=包含), i)
    for i, s in enumerate(secs):
        st = _norm(s.get("title", ""))
        if st == target:
            exact = 0
        elif target and target in st:
            exact = 1
        else:
            continue
        anc_have = _ancestor_titles(secs, i)
        # 每个想要的祖先段：只要在祖先链里被"包含"即算命中（宽松匹配层级标题差异）
        hits = 0
        for w in anc_wanted:
            if any(w == a or w in a or a in w for a in anc_have):
                hits += 1
        cands.append((hits, exact, i))
    if not cands:
        return None
    # 祖先命中越多越好；精确标题优先；同分取文档中靠前者
    cands.sort(key=lambda t: (-t[0], t[1], t[2]))
    return cands[0][2]


def find_section(refdoc, spec):
    """按 spec 找一节，返回下标或 None。
    spec 可为：层级编号(1.1.1，本方案主键)、原生编号(3.2)、纯标题名，
    或点号路径"祖先.….目标"（用于消歧重名标题）。"""
    secs = refdoc.get("sections", [])
    if not spec:
        return None
    spec_s = str(spec).strip()
    # ⓪ 纯层级编号（1 / 1.1 / 1.1.1）：用与 show 同源的 _hier_numbers 精确命中——绑定主键。
    #    这是唯一确定、可复现、天生消歧（编号唯一）的定位方式；命中即返回，不再走后续模糊匹配。
    if re.fullmatch(r'\d+(?:\.\d+)*', spec_s):
        nums = _hier_numbers(secs)
        for i, n in enumerate(nums):
            if n == spec_s:
                return i
        # 纯数字号没对上：多半是号写错或文件已改动 → 交给下面的原生编号/标题兜底，最终没有则 None
    spec_n = _norm(spec)
    # ① 整体先按原始行的编号前缀精确/前缀匹配，或标题全等（"3.2" 命中 "3.2 数据治理"）
    for i, s in enumerate(secs):
        raw_n = _norm(s.get("raw", ""))
        if raw_n.startswith(spec_n) or _norm(s.get("title", "")) == spec_n:
            return i
    # ② 点号路径"祖先.….目标"：仅当拆出多段且含非数字段时启用（纯数字编号不进这里）
    segs = [x.strip() for x in re.split(r'[.．。]', spec) if x.strip()]
    if len(segs) > 1 and any(re.search(r'[^\d]', x) for x in segs):
        idx = _find_by_path(secs, segs)
        if idx is not None:
            return idx
        spec_n = _norm(segs[-1])   # 路径没命中时回退：仅用末段做包含匹配
    # ③ 标题包含
    for i, s in enumerate(secs):
        if spec_n and spec_n in _norm(s.get("title", "")):
            return i
    return None


def slice_section(refdoc, spec, cap_chars=DEFAULT_CAP_CHARS):
    """确定性取一节（含子节）正文，按 cap 截断。找不到返回 ""。"""
    idx = find_section(refdoc, spec)
    if idx is None:
        return ""
    return _truncate(_section_scope_text(refdoc["sections"], idx), cap_chars)


def _score(title, keywords, section):
    """标题相似度 + 关键词命中率的加权分。"""
    st = _norm(section.get("title", ""))
    body = section.get("title", "") + section.get("text", "")
    sim = difflib.SequenceMatcher(None, _norm(title), st).ratio()
    kws = [k for k in (keywords or []) if k]
    hit = (sum(1 for k in kws if k in body) / len(kws)) if kws else 0.0
    return sim * 0.6 + hit * 0.4


def auto_match(refdoc, title, keywords, topn=3):
    """关键词/标题相似度排序，返回 [(section, score), ...]（降序，取前 topn）。"""
    scored = [(_score(title, keywords, s), i) for i, s in enumerate(refdoc.get("sections", []))]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [(refdoc["sections"][i], sc) for sc, i in scored[:topn]]


# ----------------------------- 生成期：解析 ref 继承 -----------------------------
def _resolve_path(fp, base_dir):
    if not fp:
        return fp
    return fp if os.path.isabs(fp) else os.path.normpath(os.path.join(base_dir, fp))


def resolve_refs(plan, base_dir, cap_chars=DEFAULT_CAP_CHARS):
    """遍历 plan.outline，按继承规则给每个叶子挂 _ref，返回引用清单 usage。
      继承：某节点 ref.file 覆盖继承；ref.exclude=True 切断（事前否决）；否则沿用父。
      叶子：section 指定 → 确定性切片；未指定 → 相似度自动匹配最优节；
            匹配分低于阈值或切不到文本 → 不挂 _ref（退回纯大模型生成，兜底）。
    _ref = {"file":显示名, "section":实际用的节标题, "text":素材, "auto":bool}
    usage 每项 = {"leaf":标题, "file":..., "section":..., "auto":bool, "score":..., "chars":..}
    对主流程零侵入：outline 中无任何 ref 时本函数早退、什么都不做。"""
    outline = plan.get("outline", [])
    if not _any_ref(outline):
        return []
    metrics = plan.get("dictionary", {}).get("metrics", []) or []   # verbatim 指标覆盖用
    docs = {}          # abspath -> RefDoc（解析缓存）

    def load(fp):
        ap = _resolve_path(fp, base_dir)
        if ap not in docs:
            docs[ap] = parse_reference(ap)
        return docs[ap]

    usage = []

    def walk(nodes, inherited, depth):
        # inherited = 来自祖先的 ref（已生效，供本层子叶继承）或 None
        for nd in nodes:
            cur = inherited
            is_own = False               # 本叶是否有"自己的"显式绑定（区别于继承）
            own = nd.get("ref")
            if isinstance(own, dict):
                if own.get("exclude"):
                    cur = None            # 事前否决：切断继承
                elif own.get("file"):
                    cur, is_own = own, True
            elif isinstance(own, list) and any(
                    isinstance(x, dict) and x.get("file") for x in own):
                cur, is_own = own, True   # 多参考绑定（list）：本叶精准绑定，合并注入
            if is_leaf(nd):
                if cur and _ref_has_file(cur):
                    _attach_leaf(nd, cur, is_own, load, cap_chars, usage, metrics, depth)
            elif nd.get("children"):
                walk(nd["children"], cur, depth + 1)

    walk(outline, None, 0)
    return usage


def _scope_indices(sections, idx):
    """第 idx 节及其所有更深层子节的下标（到出现同级/更浅标题为止）。"""
    base = sections[idx]["level"]
    out = [idx]
    for j in range(idx + 1, len(sections)):
        if sections[j]["level"] <= base:
            break
        out.append(j)
    return out


def _best_within(doc, indices, title, keywords):
    """在给定候选节下标里按相似度选最相关的一节，返回 (idx, score)。"""
    scored = sorted(((_score(title, keywords, doc["sections"][i]), i) for i in indices),
                    key=lambda x: x[0], reverse=True)
    return (scored[0][1], scored[0][0]) if scored else (indices[0], 0.0)


# 数值 token：可含比较/约束前缀 + 数字 + 常见单位；用于照抄模式的指标就近替换
_VAL_TOKEN = (r'(?:[≥≤<>=约]|不低于|不超过|不少于|不大于|不小于|大于|小于|达到|高于|低于)*\s*'
              r'\d+(?:\.\d+)?\s*'
              r'(?:%|万元|万|千|亿|毫秒|秒|分钟|小时|天|条|类|个|项|次|人|路|核|'
              r'TPS|QPS|ms|GB|TB|MB)?')


def _num_of(s):
    m = re.search(r'\d+(?:\.\d+)?', s or "")
    return m.group(0) if m else None


def apply_metric_overrides(text, metrics):
    """把"指标名 + 就近数值"里与数据字典冲突的数值替换为字典值（完整照抄模式的兜底）。
    只处理"名称紧跟数值"这一保守情形，误伤低；返回 (新文本, [(名称,原值,字典值), ...])。"""
    replaced = []
    for m in metrics or []:
        name = (m.get("name") or "").strip()
        val = (m.get("value") or "").strip()
        if not name or not val or not re.search(r'\d', val):
            continue
        dict_num = _num_of(val)
        pat = re.compile(re.escape(name) + r'\s*[:：为]?\s*(' + _VAL_TOKEN + r')')

        def _sub(mo, _name=name, _val=val, _dn=dict_num):
            cur = mo.group(1)
            cur_num = _num_of(cur)
            if cur_num and cur_num != _dn:
                replaced.append((_name, cur.strip(), _val))
                return f"{_name}{_val}"      # 名称后的数值段整体换成字典完整值
            return mo.group(0)

        text = pat.sub(_sub, text)
    return text, replaced


def _split_paras(text):
    """参考正文切成段：先按空行，再退化到按单行；去空白。"""
    out = []
    for chunk in re.split(r'\n\s*\n', (text or "").replace("\r", "")):
        chunk = chunk.strip()
        if not chunk:
            continue
        for line in chunk.split("\n"):
            s = line.strip()
            if s:
                out.append(s)
    return out


def reconstruct_blocks(doc, use_idx, leaf_depth, metrics=None):
    """把参考文件某节（含更深子节）重建为内容块 h*/p，供保结构②/完整照抄③使用。
      · 子节标题 → 更深一级 h 块（引擎后续按投标规范排版、必要时转 hnum，不保留参考原格式）；
      · 正文 → 段落 p 块；本节自身标题不重复（引擎已输出叶子标题）。
      · metrics 非空（照抄）→ 段落做"名称+就近数值"字典覆盖。
    返回 (blocks, replaced)。"""
    secs = doc["sections"]
    base = secs[use_idx]["level"]
    blocks, replaced = [], []

    def _emit(txt):
        for para in _split_paras(txt):
            if metrics:
                para, rep = apply_metric_overrides(para, metrics)
                replaced.extend(rep)
            blocks.append({"type": "p", "text": para})

    _emit(secs[use_idx].get("text", ""))
    for j in _scope_indices(secs, use_idx)[1:]:
        rel = max(secs[j]["level"] - base, 1)
        hlvl = min(leaf_depth + 1 + rel, 5)         # 叶子在 h(leaf_depth+1)，子标题须更深
        title = (secs[j].get("title") or "").strip()
        if title:
            blocks.append({"type": f"h{hlvl}", "text": title})
        _emit(secs[j].get("text", ""))
    return blocks, replaced


def _ref_has_file(ref):
    """ref（单条 dict 或多条 list）是否含可用的 file 绑定。"""
    if isinstance(ref, dict):
        return bool(ref.get("file"))
    if isinstance(ref, list):
        return any(isinstance(x, dict) and x.get("file") for x in ref)
    return False


def _fail_item(title, ref, reason):
    """精准绑定解析失败时的"响亮"记账项（fail-loud）：写进 usage 供 check_consistency 醒目告警，
    避免用户明确绑了却静默退回纯生成、事后才从成稿看出跑题。"""
    return {"leaf": title, "file": os.path.basename(str(ref.get("file", "") or "")),
            "section": ref.get("section"), "auto": False, "mode": ref.get("mode") or "adapt",
            "status": "未匹配", "reason": reason}


def _resolve_binding(ref, is_own, load, cap_chars, metrics, title, kws, leaf_depth):
    """把单条 ref 绑定解析成 (entry, usage_item)。
      成功 → (entry, usage_item)；
      精准绑定(is_own=True)解析失败 → (None, 失败记账项)  ← fail-loud，会被计入 usage 报出；
      自动/继承匹配不到 → (None, None)  ← 这是设计内的兜底纯生成，静默即可。
      is_own=True（本叶精准绑定某节）→ 确定性取整节（含子节），用 OWN_CAP_CHARS 尽量不截断。
      is_own=False（继承自父章）→ 在父绑定节范围内按相似度 narrow 到最相关子节，强制 adapt。
      section 为空（自动匹配）→ 全文相似度匹配，低于阈值则返回空（退回纯生成）。
      mode：adapt 改写借鉴；keep_headings 保子标题优化；verbatim 完整照抄。"""
    doc = load(ref["file"])
    if not doc.get("sections"):
        # 精准绑定了一份解析不出章节的文件 → 响亮报出；非精准(继承/自动)则静默兜底。
        return None, (_fail_item(title, ref, "参考文件解析不出章节") if is_own else None)
    disp = os.path.basename(doc["path"])
    mode = (ref.get("mode") or "adapt")
    if not is_own:
        mode = "adapt"          # 继承来的一律改写借鉴
    section_spec = ref.get("section")
    if section_spec:
        idx = find_section(doc, section_spec)
        if idx is None:
            # 明确指定了章节号/标题却对不上 → 最典型的静默失效点，必须响亮报出。
            return None, (_fail_item(title, ref, f"章节号/标题「{section_spec}」在参考文件里找不到对应节")
                          if is_own else None)
        if is_own:
            use_idx, score, auto = idx, 1.0, False
        else:
            use_idx, score = _best_within(doc, _scope_indices(doc["sections"], idx), title, kws)
            auto = True
    else:
        cands = auto_match(doc, title, kws, topn=1)
        if not cands or cands[0][1] < MIN_MATCH_SCORE:
            return None, None           # 参考里没有匹配内容 → 兜底纯生成（设计内，静默）
        sec = cands[0][0]
        use_idx = next(i for i, s in enumerate(doc["sections"]) if s is sec)
        score, auto = cands[0][1], True
    text = _section_scope_text(doc["sections"], use_idx)
    # adapt 也做确定性指标兜底：与字典冲突的指标在进提示词前就近替换为字典值，
    # 不再单靠模型自觉（keep_headings/verbatim 走 reconstruct_blocks 各自处理）。
    adapt_rep = []
    if mode == "adapt" and metrics:
        text, adapt_rep = apply_metric_overrides(text, metrics)
    # 精准绑定用大额上限尽量整节注入；继承/自动匹配仍按默认上限控 token。
    cap = OWN_CAP_CHARS if is_own else cap_chars
    text = _truncate(text, cap)
    if not text.strip():
        # 命中的节没有可用正文（如纯标题节）→ 精准绑定同样响亮报出，提示改绑到有正文的节。
        return None, (_fail_item(title, ref, "命中的参考节没有可用正文，建议改绑到有正文的节")
                      if is_own else None)
    sec_title = doc["sections"][use_idx].get("title", "")
    entry = {"file": disp, "section": sec_title, "text": text, "auto": auto, "mode": mode}
    u = {"leaf": title, "file": disp, "section": sec_title, "auto": auto, "mode": mode,
         "score": round(score, 3), "chars": len(text)}
    if mode == "adapt" and adapt_rep:
        u["replaced"] = len(adapt_rep)
    if mode in ("keep_headings", "verbatim"):
        blocks, replaced = reconstruct_blocks(
            doc, use_idx, leaf_depth,
            metrics=metrics if mode == "verbatim" else None)
        entry["blocks"] = blocks
        if mode == "verbatim" and replaced:
            entry["replaced"] = replaced
            u["replaced"] = len(replaced)
    return entry, u


def _merge_entries(entries):
    """多条精准绑定合并成一个 _ref：正文各带出处小标后拼接；若同为保结构/照抄则拼接 blocks。"""
    files = []
    for e in entries:
        if e["file"] not in files:
            files.append(e["file"])
    section = "；".join(f"《{e['file']}》「{e['section']}」" for e in entries)
    parts = [f"【摘自《{e['file']}》「{e['section']}」节】\n{e['text']}" for e in entries]
    modes = {e.get("mode", "adapt") for e in entries}
    mode = next(iter(modes)) if len(modes) == 1 else "adapt"
    merged = {"file": "、".join(files), "section": section,
              "text": "\n\n".join(parts),
              "auto": all(e.get("auto") for e in entries), "mode": mode}
    if mode in ("keep_headings", "verbatim") and all(e.get("blocks") for e in entries):
        blocks = []
        for e in entries:
            blocks.extend(e["blocks"])
        merged["blocks"] = blocks
    return merged


def _attach_leaf(nd, ref, is_own, load, cap_chars, usage, metrics=None, leaf_depth=0):
    """给叶子挂 _ref。ref 可为单条 dict，或多条 list（多参考合并注入到同一节）。
    单条解析走 _resolve_binding；多条各自解析后由 _merge_entries 合并为一个 _ref。"""
    title = nd.get("title", "")
    kws = nd.get("must_keywords", []) or []
    refs = ref if isinstance(ref, list) else [ref]
    entries, uitems = [], []
    for r in refs:
        if not (isinstance(r, dict) and r.get("file") and not r.get("exclude")):
            continue
        e, u = _resolve_binding(r, is_own, load, cap_chars, metrics, title, kws, leaf_depth)
        if e:
            entries.append(e)
            uitems.append(u)
        elif u:
            usage.append(u)          # fail-loud：精准绑定解析失败也记账，供审查环节醒目报出
    if not entries:
        return                       # 无可用素材：失败项（若有）已入 usage，本叶退回纯生成
    nd["_ref"] = entries[0] if len(entries) == 1 else _merge_entries(entries)
    usage.extend(uitems)


def _any_ref(nodes):
    for nd in nodes:
        r = nd.get("ref")
        if isinstance(r, dict) and (r.get("file") or r.get("exclude")):
            return True
        if isinstance(r, list) and any(isinstance(x, dict) and x.get("file") for x in r):
            return True
        if nd.get("children") and _any_ref(nd["children"]):
            return True
    return False


# ----------------------------- 交互绑定 CLI -----------------------------
def _index_nodes(outline):
    """返回 [(id, depth, node)]（文档顺序），id 与 plan.py assign_ids 口径一致。"""
    out = []

    def walk(nodes, prefix, depth):
        for i, nd in enumerate(nodes):
            nid = f"{prefix}{i}" if prefix == "" else f"{prefix}.{i}"
            out.append((nid, depth, nd))
            if nd.get("children"):
                walk(nd["children"], nid, depth + 1)

    walk(outline, "", 0)
    return out


def _print_refdoc(refdoc, fno):
    tag = "可靠" if refdoc.get("reliable") else "需确认(降级)"
    print(f"\n[文件{fno}] {os.path.basename(refdoc['path'])}  解析:{refdoc['kind']}  {tag}")
    for w in refdoc.get("warnings", []):
        print(f"    ! {w}")
    secs = refdoc.get("sections", [])
    nums = _hier_numbers(secs)                       # 展示号与 find_section 解析号同源
    for i, s in enumerate(secs):
        prev = (s.get("text", "").strip().replace("\n", " ")[:28])
        print(f"    {nums[i]}  {'  ' * (s['level'] - 1)}{s.get('title', '')}"
              + (f"  〔{prev}…〕" if prev else "  〔（无正文）〕"))


def _bid_hier(outline):
    """遍历嵌套大纲，产出 [(层级号, node)]，号形如 1 / 1.1 / 2.3.1，drill 到每一级。"""
    out = []

    def walk(nodes, prefix):
        for i, nd in enumerate(nodes, 1):
            num = f"{prefix}.{i}" if prefix else str(i)
            out.append((num, nd))
            if nd.get("children"):
                walk(nd["children"], num)

    walk(outline, "")
    return out


def _print_bidoutline(outline):
    print("\n===== 本投标文件大纲（层级编号，drill 到每一级叶子）=====")
    for num, nd in _bid_hier(outline):
        leaf = "  <叶子>" if is_leaf(nd) else ""
        indent = "  " * num.count(".")
        print(f"  {num}  {indent}{nd.get('title', '')}{leaf}")


def _descendant_leaf_ids(outline, node_id):
    """给定节点 id，返回其子树下全部叶子的 id 列表（含自身若自身即叶子）。"""
    ids = []
    for nid, _d, nd in _index_nodes(outline):
        if (nid == node_id or nid.startswith(node_id + ".")) and is_leaf(nd):
            ids.append(nid)
    return ids


def _bind_interactive(outline, refdocs):
    idmap = {nid: nd for nid, _d, nd in _index_nodes(outline)}
    print("\n===== 本标书大纲（可绑定任意层级节点，绑父章其下子叶自动继承）=====")
    for nid, depth, nd in _index_nodes(outline):
        leaf = "  <叶子>" if is_leaf(nd) else ""
        print(f"  id={nid}{'  ' * depth}  {nd.get('title', '')}{leaf}")
    print("\n绑定语法（每行一条，回车空行结束）：")
    print("  绑定：  <大纲id> <文件号>:<层级号>    例  0.1 1:3.1  （0.1 节参考 文件1 的 3.1 节）")
    print("  自动：  <大纲id> <文件号>            例  0.1 1     （节号交给相似度自动匹配）")
    print("  否决：  <大纲id> x                   例  0.1 x     （本节强制不参考，事前否决）")
    while True:
        line = input("  > ").strip()
        if not line:
            break
        parts = line.split()
        if len(parts) < 2:
            print("    格式不对，跳过。")
            continue
        nid = parts[0]
        if nid not in idmap:
            print(f"    没有 id={nid}，跳过。")
            continue
        nd = idmap[nid]
        if parts[1].lower() in ("x", "exclude", "排除"):
            nd["ref"] = {"exclude": True}
            print(f"    已设为不参考：{nd.get('title', '')}")
            continue
        fspec = parts[1]
        fno, _, sno = fspec.partition(":")
        try:
            fi = int(fno) - 1
            refdoc = refdocs[fi]
        except (ValueError, IndexError):
            print(f"    文件号 {fno} 无效，跳过。")
            continue
        sec_title = None
        sec_spec = None
        if sno:
            spec = sno.lstrip("#")                       # 兼容旧写法误带 #
            idx = find_section(refdoc, spec)
            if idx is None:
                print(f"    节号 {sno} 在该文件里找不到对应节，跳过。")
                continue
            sec_title = refdoc["sections"][idx].get("title", "")
            sec_spec = spec                              # 写回层级号，生成期同源解析
        nd["ref"] = {"file": os.path.basename(refdoc["path"]),
                     "section": sec_spec if sec_spec else sec_title}
        # 预览确认：绑定即回显将用到的内容片段
        preview_spec = sec_spec if sec_spec else None
        if preview_spec:
            snip = slice_section(refdoc, preview_spec, 120)
        else:
            cands = auto_match(refdoc, nd.get("title", ""),
                               nd.get("must_keywords", []) or [], topn=1)
            if cands and cands[0][1] >= MIN_MATCH_SCORE:
                snip = _truncate(cands[0][0].get("text", ""), 120)
                sec_title = cands[0][0].get("title", "")
            else:
                snip = ""
        where = f"{sno}「{sec_title}」" if sno else (f"自动→「{sec_title}」" if sec_title else "自动(暂未匹配到)")
        print(f"    绑定 {nd.get('title', '')} → 文件{fno} {where}")
        print(f"      预览：{snip[:80] or '（无正文，建议改绑到有正文的节）'}…")
        # 绑非叶子：回显其子叶，允许逐个事前否决
        if not is_leaf(nd):
            kids = _descendant_leaf_ids(outline, nid)
            print(f"      此绑定将被以下 {len(kids)} 个子叶继承：{'、'.join(kids)}")
            ex = input("      要事前排除的子叶 id（空格分隔，回车=全部继承）> ").strip()
            for xid in ex.split():
                if xid in idmap and is_leaf(idmap[xid]):
                    idmap[xid]["ref"] = {"exclude": True}
                    print(f"        已排除：{idmap[xid].get('title', '')}")
    return outline


def cmd_bind(args):
    with open(args.outline, "r", encoding="utf-8") as f:
        data = json.load(f)
    outline = data.get("outline", [])
    files = [p.strip() for p in args.refs.split(",") if p.strip()]
    base = os.path.dirname(os.path.abspath(args.outline))
    refdocs = [parse_reference(_resolve_path(p, base)) for p in files]
    for i, rd in enumerate(refdocs, 1):
        _print_refdoc(rd, i)
    if args.yes:
        print("\n[--yes] 跳过交互绑定，未写入任何 ref。")
    else:
        _bind_interactive(outline, refdocs)
    out = args.out or args.outline
    with open(out, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    n = sum(1 for _n, _d, nd in _index_nodes(outline) if isinstance(nd.get("ref"), dict))
    print(f"\n已写入 {out}（{n} 个节点带 ref 绑定）。")
    return 0


def main():
    ap = argparse.ArgumentParser(description="参考资料解析/切片/匹配/交互绑定")
    sub = ap.add_subparsers(dest="cmd")

    b = sub.add_parser("bind", help="交互式把大纲章节绑定到参考文件某节，写回 ref")
    b.add_argument("--outline", required=True)
    b.add_argument("--refs", required=True, help="参考文件路径，逗号分隔（相对 outline 目录）")
    b.add_argument("--out", default=None, help="输出（默认覆盖 outline）")
    b.add_argument("--yes", action="store_true", help="只解析预览、不交互绑定")

    s = sub.add_parser("show", help="只解析并打印参考文件的章节树")
    s.add_argument("--refs", required=True)

    o = sub.add_parser("outline", help="打印投标大纲的层级编号树（drill 到每一级叶子）")
    o.add_argument("--outline", required=True)

    p = sub.add_parser("pair", help="同时打印投标大纲与参考文件章节树（都带层级编号），供人工指定绑定")
    p.add_argument("--outline", required=True)
    p.add_argument("--refs", required=True, help="参考文件路径，逗号分隔（相对 outline 目录）")

    args = ap.parse_args()
    if args.cmd == "bind":
        return cmd_bind(args)
    if args.cmd == "show":
        for i, p in enumerate([x.strip() for x in args.refs.split(",") if x.strip()], 1):
            _print_refdoc(parse_reference(p), i)
        return 0
    if args.cmd == "outline":
        with open(args.outline, "r", encoding="utf-8") as f:
            data = json.load(f)
        _print_bidoutline(data.get("outline", []))
        return 0
    if args.cmd == "pair":
        with open(args.outline, "r", encoding="utf-8") as f:
            data = json.load(f)
        _print_bidoutline(data.get("outline", []))
        base = os.path.dirname(os.path.abspath(args.outline))
        for i, rp in enumerate([x.strip() for x in args.refs.split(",") if x.strip()], 1):
            _print_refdoc(parse_reference(_resolve_path(rp, base)), i)
        return 0
    ap.print_help()
    return 1


if __name__ == "__main__":
    sys.exit(main())
